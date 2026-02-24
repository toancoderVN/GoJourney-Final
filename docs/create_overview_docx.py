from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os

def set_cell_border(cell, **kwargs):
    """
    Set cell`s border
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    for edge in ('top', 'start', 'bottom', 'end', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = tcPr.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcPr.append(element)
            
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

def create_overview_document():
    doc = Document()

    # --- STYLE CONFIGURATION ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # 1. Main Title
    title = doc.add_paragraph('BẢNG TỔNG QUAN CHỨC NĂNG HỆ THỐNG')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_run.font.color.rgb = RGBColor(44, 62, 80)
    doc.add_paragraph('')

    # --- DATA PREPARATION (EXACTLY AS REQUESTED) ---
    # Column: [STT, Tên UseCase, Đối tượng, Mô tả]
    data = [
        ("3.3.1", "UseCase Đăng ký tài khoản mới", "Khách (Guest)", "Đăng ký tài khoản mới vào hệ thống sử dụng Email và Mật khẩu."),
        ("3.3.2", "UseCase Đăng nhập", "Người dùng", "Đăng nhập vào hệ thống để truy cập các tính năng."),
        ("3.3.3", "UseCase Đăng xuất", "Người dùng", "Đăng xuất khỏi hệ thống, kết thúc phiên làm việc."),
        ("3.3.4", "UseCase Xem thông tin profile người dùng", "Người dùng", "Xem và quản lý thông tin cá nhân của tài khoản."),
        ("3.3.5", "UseCase Xem danh sách người đồng hành", "Người dùng", "Xem danh sách các người đồng hành (bạn bè) đã kết nối."),
        ("3.3.6", "UseCase Xem pending invitations", "Người dùng", "Xem danh sách các lời mời kết nối đang chờ xử lý."),
        ("3.3.7", "UseCase Connect với user qua code", "Người dùng", "Kết nối với người dùng khác thông qua mã định danh (Code)."),
        ("3.3.8", "UseCase Tạo trip từ AI agent booking", "Người dùng", "Chuyển đổi thông tin đặt phòng thành công từ AI Agent thành một chuyến đi."),
        ("3.3.9", "UseCase Đàm phán với khách sạn qua Zalo (tự động)", "System (AI)", "AI Agent tự động gửi tin nhắn Zalo để đàm phán giá và phòng với khách sạn."),
        ("3.3.10", "UseCase Xác nhận thanh toán booking", "Người dùng", "Người dùng xác nhận đồng ý thanh toán/đặt cọc theo yêu cầu từ khách sạn."),
        ("3.3.11", "UseCase Hủy booking request", "Người dùng", "Hủy bỏ yêu cầu đặt phòng đang xử lý."),
        ("3.3.12", "UseCase Kết nối tài khoản Zalo", "Người dùng", "Liên kết tài khoản Zalo cá nhân vào hệ thống để AI sử dụng."),
        ("3.3.13", "UseCase Xem danh sách Zalo conversations", "Người dùng", "Xem danh sách các cuộc hội thoại Zalo đã đồng bộ."),
        ("3.3.14", "UseCase Gửi tin nhắn qua Zalo tới khách sạn", "System/AI", "Gửi tin nhắn văn bản tới tài khoản Zalo của khách sạn."),
        ("3.3.15", "UseCase Nhận tin nhắn phản hồi từ khách sạn", "System", "Hệ thống nhận và xử lý tin nhắn phản hồi từ khách sạn qua Zalo Webhook."),
        ("3.3.16", "UseCase Chat với AI assistant", "Người dùng", "Trò chuyện trực tiếp với trợ lý ảo AI để được hỗ trợ."),
        ("3.3.17", "UseCase Nghiên cứu sâu về điểm đến (Deep Research)", "Người dùng", "Yêu cầu AI nghiên cứu chi tiết về một địa điểm du lịch cụ thể."),
        ("3.3.18", "UseCase Tìm kiếm thông tin du lịch trên web", "System (AI)", "AI tự động tìm kiếm thông tin du lịch mới nhất trên internet."),
        ("3.3.19", "UseCase Mời companion tham gia chuyến đi", "Người dùng", "Mời người đồng hành (đã kết nối) tham gia vào một chuyến đi cụ thể."),
        ("3.3.20", "UseCase Kết nối với người dùng khác qua User ID", "Người dùng", "Gửi yêu cầu kết bạn với người dùng khác thông qua User ID."),
        ("3.3.21", "UseCase Xem danh sách lời mời chờ xử lý", "Người dùng", "Kiểm tra và quản lý các lời mời kết nối/tham gia chuyến đi đang chờ."),
        ("3.3.22", "UseCase Xóa người đồng hành", "Người dùng", "Hủy kết nối (xóa) một người dùng khỏi danh sách người đồng hành.")
    ]

    # --- TABLE CREATION ---
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    
    # Header Row
    hdr_cells = table.rows[0].cells
    headers = ['STT', 'Tên Use Case', 'Đối tượng', 'Mô tả ngắn']
    
    for i, header_text in enumerate(headers):
        cell = hdr_cells[i]
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(header_text)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Blue background
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '4472C4')
        tcPr.append(shd)

    # Data Rows
    for stt, name, actor, desc in data:
        row_cells = table.add_row().cells
        
        # STT
        p0 = row_cells[0].paragraphs[0]
        p0.add_run(stt).bold = True
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Name
        p1 = row_cells[1].paragraphs[0]
        p1.add_run(name)
        
        # Actor
        p2 = row_cells[2].paragraphs[0]
        p2.add_run(actor)
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Desc
        p3 = row_cells[3].paragraphs[0]
        p3.add_run(desc)

        # Borders
        for cell in row_cells:
            set_cell_border(
                cell,
                top={"sz": 4, "val": "single", "color": "auto"},
                bottom={"sz": 4, "val": "single", "color": "auto"},
                start={"sz": 4, "val": "single", "color": "auto"},
                end={"sz": 4, "val": "single", "color": "auto"},
            )

    # Column widths
    for row in table.rows:
        row.cells[0].width = Inches(0.8)  # STT
        row.cells[1].width = Inches(3.0)  # Name
        row.cells[2].width = Inches(1.2)  # Actor
        row.cells[3].width = Inches(3.5)  # Desc

    output_path = os.path.join(os.getcwd(), 'system_functions_overview.docx')
    doc.save(output_path)
    print(f"File created successfully at: {output_path}")

if __name__ == "__main__":
    create_overview_document()
