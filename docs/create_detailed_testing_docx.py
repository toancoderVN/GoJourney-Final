from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for edge in ('top', 'start', 'bottom', 'end', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = tcPr.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcPr.append(element)
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

def create_doc():
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)

    # Title
    title = doc.add_paragraph('5.3. CHI TIẾT KỊCH BẢN KIỂM THỬ CHỨC NĂNG')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.bold = True
    title_run.font.size = Pt(16)
    doc.add_paragraph('')

    md_path = 'docs/detailed_functional_testing.md'
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    # Split by Table Headers (### Bảng...)
    sections = re.split(r'(### Bảng 5\.1\.\d+: [^\n]+)', content)
    
    for section in sections:
        section = section.strip()
        if not section: continue
        
        if section.startswith('### '):
            # Table Title
            p = doc.add_paragraph(section.replace('### ', '').strip())
            p.style = 'Heading 3'
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            
        elif '|' in section and 'STT' in section:
            # Table Content
            lines = section.split('\n')
            # Extract header and rows
            rows = []
            for line in lines:
                if '|' in line and '---' not in line:
                    cols = [c.strip() for c in line.split('|') if c.strip()]
                    if cols:
                        rows.append(cols)
            
            if not rows: continue

            # Create Doc Table
            headers = rows[0]
            data_rows = rows[1:]
            
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = 'Table Grid'
            table.autofit = False 
            
            # Set Header
            hdr_cells = table.rows[0].cells
            for i, h in enumerate(headers):
                p = hdr_cells[i].paragraphs[0]
                run = p.add_run(h)
                run.bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Gray Background
                tcPr = hdr_cells[i]._tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'E7E6E6')
                tcPr.append(shd)

            # Set Data
            for r_data in data_rows:
                row = table.add_row()
                for i, val in enumerate(r_data):
                    if i < len(row.cells):
                        # Handle line breaks <br> from markdown
                        cell_text = val.replace('<br>', '\n').replace('**', '')
                        row.cells[i].text = cell_text
            
            # Formatting Column Widths (Approximate)
            # STT: Small
            table.columns[0].width = Inches(0.4)
            # Tên: Medium
            table.columns[1].width = Inches(1.5)
            # Bước: Large
            table.columns[2].width = Inches(2.5)
            # Expect: Medium-Large
            table.columns[3].width = Inches(2.0)
            
            doc.add_paragraph('') # Spacing after table

    doc.save('docs/detailed_functional_testing.docx')
    print("Created docs/detailed_functional_testing.docx")

if __name__ == "__main__":
    create_doc()
