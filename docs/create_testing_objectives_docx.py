from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_objectives_doc():
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)

    md_path = 'docs/testing_objectives.md'
    if not os.path.exists(md_path):
        print(f"File {md_path} not found")
        return

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('# '):
            # Heading 1
            p = doc.add_paragraph(line.replace('# ', '').upper())
            p.style = 'Heading 1' # Use Word's generic heading style if available, or format manually
            run = p.runs[0]
            run.bold = True
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0, 0, 0)
            p.paragraph_format.space_after = Pt(12)
            
        elif line.startswith('### '):
            # Heading 3 (Objectives Title)
            p = doc.add_paragraph(line.replace('### ', ''))
            run = p.runs[0]
            run.bold = True
            run.italic = True
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            
        elif line.startswith('**'):
            # Bold Categories (e.g., **1. Kiểm thử tính khả dụng:**)
            p = doc.add_paragraph(line.replace('**', ''))
            p.runs[0].bold = True
            p.paragraph_format.space_before = Pt(6)
            
        elif line.startswith('*   '):
            # Bullet points
            p = doc.add_paragraph(line.replace('*   ', ''), style='List Bullet')
            
        else:
            # Normal text
            doc.add_paragraph(line)

    output_path = 'docs/testing_objectives.docx'
    doc.save(output_path)
    print(f"Created {output_path}")

if __name__ == "__main__":
    create_objectives_doc()
