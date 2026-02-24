#!/usr/bin/env python3
"""Convert Trip Invitation Use Cases to DOCX format - Single table format"""

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def merge_cells(table, row_idx, col_start, col_end):
    """Merge cells in a row"""
    cell = table.cell(row_idx, col_start)
    for i in range(col_start + 1, col_end + 1):
        cell.merge(table.cell(row_idx, i))
    return cell

def add_usecase(doc, title, data):
    """Add a use case with single table format"""
    doc.add_heading(title, level=2)
    
    table = doc.add_table(rows=0, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    for item in data:
        row_type = item.get('type', 'single')
        
        if row_type == 'single':
            row = table.add_row()
            row.cells[0].text = item['label']
            row.cells[0].paragraphs[0].runs[0].bold = True
            
            merged = merge_cells(table, len(table.rows) - 1, 1, 2)
            merged.text = item['value']
            
        elif row_type == 'header_flow':
            row = table.add_row()
            row.cells[0].text = item['label']
            row.cells[0].paragraphs[0].runs[0].bold = True
            
            row.cells[1].text = 'Actor'
            row.cells[1].paragraphs[0].runs[0].bold = True
            
            row.cells[2].text = 'System'
            row.cells[2].paragraphs[0].runs[0].bold = True
            
        elif row_type == 'flow':
            row = table.add_row()
            row.cells[0].text = ''
            row.cells[1].text = item.get('actor', '')
            row.cells[2].text = item.get('system', '')
            
        elif row_type == 'exception_header':
            row = table.add_row()
            row.cells[0].text = item['label']
            row.cells[0].paragraphs[0].runs[0].bold = True
            
            merged = merge_cells(table, len(table.rows) - 1, 1, 2)
            merged.text = item['value']
            
        elif row_type == 'exception':
            row = table.add_row()
            row.cells[0].text = ''
            
            merged = merge_cells(table, len(table.rows) - 1, 1, 2)
            merged.text = item['value']
    
    doc.add_paragraph()

def create_invitation_usecases_doc():
    doc = Document()
    
    # Title
    title = doc.add_heading('USE CASES - TRIP INVITATIONS', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('3.4.6. Module Quản Lý Lời Mời & Chia Sẻ Chuyến Đi', level=1)
    
    # ==========================================
    # UC 3.4.6.1 - Mời companion tham gia chuyến đi
    # ==========================================
    add_usecase(doc, '3.4.6.1. UseCase Mời companion tham gia chuyến đi', [
        {'type': 'single', 'label': 'Usecase', 'value': 'Mời companion tham gia chuyến đi'},
        {'type': 'single', 'label': 'Ngữ cảnh', 'value': 'Người tạo chuyến đi (Trip Owner) muốn mời bạn bè, người thân cùng tham gia vào kế hoạch'},
        {'type': 'single', 'label': 'Mô tả', 'value': 'Gửi lời mời tới người dùng khác có trong danh sách bạn đồng hành hoặc qua email/link'},
        {'type': 'single', 'label': 'Tác Nhân', 'value': 'Trip Owner'},
        {'type': 'single', 'label': 'Sự kiện', 'value': 'Trip Owner nhấn "Mời bạn bè" trong màn hình chi tiết chuyến đi'},
        {'type': 'single', 'label': 'Kết quả', 'value': 'Lời mời được gửi đi và hiển thị trạng thái "Pending"'},
        {'type': 'header_flow', 'label': 'Luồng sự kiện'},
        {'type': 'flow', 'actor': 'Trip Owner nhập tên/email bạn bè hoặc chọn từ danh sách Companion', 'system': 'Hệ thống hiển thị danh sách gợi ý'},
        {'type': 'flow', 'actor': 'Trip Owner nhấn "Gửi lời mời"', 'system': 'Hệ thống tạo record TripInvitation trong database'},
        {'type': 'flow', 'actor': '', 'system': 'Hệ thống gửi notification (in-app/email) tới người được mời'},
        {'type': 'flow', 'actor': '', 'system': 'Hệ thống thông báo "Đã gửi lời mời thành công"'},
        {'type': 'exception_header', 'label': 'Ngoại lệ', 'value': 'Người dùng đã được mời: Hệ thống báo "Người này đã có trong danh sách khách mời"'},
        {'type': 'exception', 'value': 'Người dùng không tồn tại: Hệ thống báo "Không tìm thấy người dùng này"'},
    ])
    
    # ==========================================
    # UC 3.4.6.2 - Chấp nhận lời mời tham gia chuyến đi
    # ==========================================
    add_usecase(doc, '3.4.6.2. UseCase Chấp nhận lời mời tham gia chuyến đi', [
        {'type': 'single', 'label': 'Usecase', 'value': 'Chấp nhận lời mời tham gia chuyến đi'},
        {'type': 'single', 'label': 'Ngữ cảnh', 'value': 'Người dùng nhận được lời mời tham gia vào một chuyến đi do người khác tổ chức'},
        {'type': 'single', 'label': 'Mô tả', 'value': 'Người dùng xem chi tiết lời mời và quyết định tham gia, khi đó họ sẽ thấy chuyến đi trong danh sách "Trips" của mình'},
        {'type': 'single', 'label': 'Tác Nhân', 'value': 'Người được mời (Invitee)'},
        {'type': 'single', 'label': 'Sự kiện', 'value': 'Người dùng nhận được thông báo lời mời và nhấn vào xem'},
        {'type': 'single', 'label': 'Kết quả', 'value': 'Người dùng trở thành thành viên của chuyến đi (Trip Member)'},
        {'type': 'header_flow', 'label': 'Luồng sự kiện'},
        {'type': 'flow', 'actor': 'Người dùng mở thông báo lời mời', 'system': 'Hệ thống hiển thị thông tin tóm tắt chuyến đi (Tên, Địa điểm, Thời gian)'},
        {'type': 'flow', 'actor': 'Người dùng nhấn "Tham gia chuyến đi"', 'system': 'Hệ thống cập nhật trạng thái invitation thành "ACCEPTED"'},
        {'type': 'flow', 'actor': '', 'system': 'Hệ thống thêm người dùng vào danh sách thành viên chuyến đi'},
        {'type': 'flow', 'actor': '', 'system': 'Hệ thống gửi thông báo cho Trip Owner: "[Tên] đã chấp nhận lời mời"'},
        {'type': 'flow', 'actor': '', 'system': 'Hệ thống chuyển hướng người dùng tới trang chi tiết chuyến đi'},
        {'type': 'exception_header', 'label': 'Ngoại lệ', 'value': 'Lời mời hết hạn/bị hủy: Hệ thống báo "Lời mời này không còn hiệu lực"'},
    ])
    
    # ==========================================
    # UC 3.4.6.3 - Xem mã cá nhân của tôi
    # ==========================================
    add_usecase(doc, '3.4.6.3. UseCase Xem mã cá nhân của tôi', [
        {'type': 'single', 'label': 'Usecase', 'value': 'Xem mã cá nhân của tôi'},
        {'type': 'single', 'label': 'Ngữ cảnh', 'value': 'Người dùng muốn lấy mã định danh duy nhất (User Code) để chia sẻ cho bạn bè kết bạn nhanh'},
        {'type': 'single', 'label': 'Mô tả', 'value': 'Hiển thị mã User Code (dạng chữ hoặc QR Code) để người khác có thể tìm kiếm và kết bạn'},
        {'type': 'single', 'label': 'Tác Nhân', 'value': 'Người dùng đã đăng nhập'},
        {'type': 'single', 'label': 'Sự kiện', 'value': 'Người dùng truy cập trang Profile hoặc Settings'},
        {'type': 'single', 'label': 'Kết quả', 'value': 'Mã cá nhân được hiển thị rõ ràng'},
        {'type': 'header_flow', 'label': 'Luồng sự kiện'},
        {'type': 'flow', 'actor': 'Người dùng chọn mục "Mã cá nhân" trong Profile', 'system': 'Hệ thống gọi API GET /travel-companions/my-code'},
        {'type': 'flow', 'actor': '', 'system': 'Hệ thống lấy mã User Code từ database'},
        {'type': 'flow', 'actor': '', 'system': 'Hệ thống hiển thị mã (VD: NAV-8392) và nút "Copy"'},
        {'type': 'flow', 'actor': 'Người dùng nhấn "Copy"', 'system': 'Hệ thống sao chép mã vào clipboard và thông báo "Đã copy"'},
        {'type': 'exception_header', 'label': 'Ngoại lệ', 'value': 'Lỗi server: Hệ thống báo "Không thể lấy mã cá nhân lúc này"'},
    ])
    
    # Save document
    output_path = '/home/navin/Working/Booking_Agents/travel-agent-mvp/docs/usecases/trip_invitation_usecases.docx'
    doc.save(output_path)
    print(f"✅ Document saved to: {output_path}")

if __name__ == "__main__":
    create_invitation_usecases_doc()
