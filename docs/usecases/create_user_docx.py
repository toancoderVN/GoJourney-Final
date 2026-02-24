#!/usr/bin/env python3
"""Convert User Use Cases to DOCX format - Single table format"""

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def merge_cells(table, row_idx, col_start, col_end):
    """Merge cells in a row"""
    cell = table.cell(row_idx, col_start)
    for i in range(col_start + 1, col_end + 1):
        cell.merge(table.cell(row_idx, i))
    return cell

def add_usecase(doc, title, data):
    """Add a use case with single table format"""
    doc.add_heading(title, level=2)
    
    # Create table with 3 columns (Label | Actor | System)
    # Column 0: Label (merged across for single rows)
    # Column 1: Actor
    # Column 2: System
    table = doc.add_table(rows=0, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    # Set column widths
    table.allow_autofit = False
    # This acts as a hint, actual width setting requires xml manipulation or consistent cell width setting
    
    for item in data:
        row_type = item.get('type', 'single')
        
        if row_type == 'single':
            # Single row: Label | Value (merged cols 1-2)
            row = table.add_row()
            row.cells[0].text = item['label']
            row.cells[0].paragraphs[0].runs[0].bold = True
            
            merged = merge_cells(table, len(table.rows) - 1, 1, 2)
            merged.text = item['value']
            
        elif row_type == 'header_flow':
            # Header for Flow: "Luồng sự kiện" | "Actor" | "System"
            row = table.add_row()
            row.cells[0].text = item['label']
            row.cells[0].paragraphs[0].runs[0].bold = True
            
            row.cells[1].text = 'Actor'
            row.cells[1].paragraphs[0].runs[0].bold = True
            
            row.cells[2].text = 'System'
            row.cells[2].paragraphs[0].runs[0].bold = True
            
        elif row_type == 'flow':
            # Flow row: Empty | Actor Action | System Action
            row = table.add_row()
            row.cells[0].text = ''
            row.cells[1].text = item.get('actor', '')
            row.cells[2].text = item.get('system', '')
            
        elif row_type == 'exception_header':
            # Exception Header (if needed distinct) or just use 'single' logic if merging
            row = table.add_row()
            row.cells[0].text = item['label']
            row.cells[0].paragraphs[0].runs[0].bold = True
            
            merged = merge_cells(table, len(table.rows) - 1, 1, 2)
            merged.text = item['value']
            
        elif row_type == 'exception':
            # Exception continuation
            row = table.add_row()
            row.cells[0].text = ''
            
            merged = merge_cells(table, len(table.rows) - 1, 1, 2)
            merged.text = item['value']
    
    doc.add_paragraph()

def create_user_usecases_doc():
    doc = Document()
    
    # Title
    title = doc.add_heading('USE CASES - USER MANAGEMENT SERVICE', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('3.4.2. Module Quản Lý Người Dùng & Bạn Đồng Hành', level=1)
    
    # ==========================================
    # UC 3.4.2.1 - Xem danh sách người đồng hành
    # ==========================================
    add_usecase(doc, '3.4.2.1. UseCase Xem danh sách người đồng hành', [
        {'type': 'single', 'label': 'Usecase', 'value': 'Xem danh sách người đồng hành'},
        {'type': 'single', 'label': 'Ngữ cảnh', 'value': 'Người dùng muốn xem danh sách bạn bè, người thân đã kết nối trong hệ thống'},
        {'type': 'single', 'label': 'Mô tả', 'value': 'Hiển thị danh sách tất cả người đồng hành đã kết nối, cùng với trạng thái và thông tin cơ bản'},
        {'type': 'single', 'label': 'Tác Nhân', 'value': 'Người dùng đã đăng nhập'},
        {'type': 'single', 'label': 'Sự kiện', 'value': 'Người dùng truy cập vào trang "Bạn đồng hành" (Travel Companions)'},
        {'type': 'single', 'label': 'Kết quả', 'value': 'Danh sách người đồng hành được hiển thị'},
        {'type': 'header_flow', 'label': 'Luồng sự kiện'},
        {'type': 'flow', 'actor': 'Người dùng chọn menu "Bạn đồng hành"', 'system': 'Hệ thống gọi API GET /travel-companions'},
        {'type': 'flow', 'actor': '', 'system': 'Hệ thống truy vấn danh sách companion từ database'},
        {'type': 'flow', 'actor': '', 'system': 'Hệ thống hiển thị danh sách dạng lưới hoặc danh sách bao gồm: Avatar, Tên, Mối quan hệ, Trạng thái (Online/Offline)'},
        {'type': 'exception_header', 'label': 'Ngoại lệ', 'value': 'Lỗi server: Hệ thống báo lỗi "Không thể tải danh sách, vui lòng thử lại sau"'},
        {'type': 'exception', 'value': 'Danh sách trống: Hệ thống hiển thị thông báo "Bạn chưa có người đồng hành nào" và gợi ý thêm mới'},
    ])
    
    # ==========================================
    # UC 3.4.2.2 - Xem pending invitations
    # ==========================================
    add_usecase(doc, '3.4.2.2. UseCase Xem pending invitations', [
        {'type': 'single', 'label': 'Usecase', 'value': 'Xem pending invitations (Lời mời đang chờ)'},
        {'type': 'single', 'label': 'Ngữ cảnh', 'value': 'Người dùng muốn kiểm tra các lời mời kết bạn chưa được xử lý'},
        {'type': 'single', 'label': 'Mô tả', 'value': 'Xem danh sách các lời mời kết bạn đã nhận nhưng chưa chấp nhận hoặc từ chối'},
        {'type': 'single', 'label': 'Tác Nhân', 'value': 'Người dùng đã đăng nhập'},
        {'type': 'single', 'label': 'Sự kiện', 'value': 'Người dùng chọn tab "Lời mời kết bạn" hoặc mở thông báo lời mời'},
        {'type': 'single', 'label': 'Kết quả', 'value': 'Hiển thị danh sách các lời mời đang chờ xử lý'},
        {'type': 'header_flow', 'label': 'Luồng sự kiện'},
        {'type': 'flow', 'actor': 'Người dùng chọn tab "Lời mời"', 'system': 'Hệ thống gọi API GET /travel-companions/invitations'},
        {'type': 'flow', 'actor': '', 'system': 'Hệ thống lọc các lời mời có trạng thái \'pending\' gửi đến user hiện tại'},
        {'type': 'flow', 'actor': '', 'system': 'Hệ thống hiển thị danh sách lời mời bao gồm: Người gửi, Thời gian gửi, Tin nhắn đính kèm'},
        {'type': 'flow', 'actor': '', 'system': 'Hiển thị 2 nút hành động: "Chấp nhận" và "Từ chối"'},
        {'type': 'exception_header', 'label': 'Ngoại lệ', 'value': 'Không có lời mời nào: Hệ thống hiển thị "Không có lời mời nào đang chờ"'},
    ])
    
    # ==========================================
    # UC 3.4.2.3 - Connect với user qua code
    # ==========================================
    add_usecase(doc, '3.4.2.3. UseCase Connect với user qua code', [
        {'type': 'single', 'label': 'Usecase', 'value': 'Connect với user qua code'},
        {'type': 'single', 'label': 'Ngữ cảnh', 'value': 'Người dùng muốn kết bạn nhanh với người khác thông qua mã kết bạn (User Code)'},
        {'type': 'single', 'label': 'Mô tả', 'value': 'Người dùng nhập mã kết bạn (ví dụ: NAV-1234) của người khác để gửi lời mời kết bạn'},
        {'type': 'single', 'label': 'Tác Nhân', 'value': 'Người dùng đã đăng nhập'},
        {'type': 'single', 'label': 'Sự kiện', 'value': 'Người dùng nhập mã vào ô "Nhập mã kết bạn" và nhấn Gửi'},
        {'type': 'single', 'label': 'Kết quả', 'value': 'Lời mời kết bạn được gửi đến người sở hữu mã'},
        {'type': 'header_flow', 'label': 'Luồng sự kiện'},
        {'type': 'flow', 'actor': 'Người dùng nhập User Code của bạn bè', 'system': 'Hệ thống hiển thị form nhập mã'},
        {'type': 'flow', 'actor': 'Người dùng chọn mối quan hệ (Bạn bè/Gia đình)', 'system': ''},
        {'type': 'flow', 'actor': 'Người dùng nhấn nút "Kết nối"', 'system': 'Hệ thống gọi API POST /travel-companions/connect-by-code'},
        {'type': 'flow', 'actor': '', 'system': 'Hệ thống kiểm tra mã code có tồn tại và hợp lệ không'},
        {'type': 'flow', 'actor': '', 'system': 'Hệ thống tạo lời mời kết bạn mới trong database'},
        {'type': 'flow', 'actor': '', 'system': 'Hệ thống gửi thông báo realtime đến người nhận'},
        {'type': 'flow', 'actor': '', 'system': 'Hệ thống thông báo "Đã gửi lời mời kết bạn thành công"'},
        {'type': 'exception_header', 'label': 'Ngoại lệ', 'value': 'Mã code không tồn tại: Hệ thống báo lỗi "Mã người dùng không tìm thấy"'},
        {'type': 'exception', 'value': 'Đã là bạn bè: Hệ thống báo lỗi "Người này đã có trong danh sách bạn đồng hành"'},
        {'type': 'exception', 'value': 'Tự kết bạn với chính mình: Hệ thống báo lỗi "Không thể kết bạn với chính mình"'},
    ])
    
    # Save document
    output_path = '/home/navin/Working/Booking_Agents/travel-agent-mvp/docs/usecases/user_usecases.docx'
    doc.save(output_path)
    print(f"✅ Document saved to: {output_path}")

if __name__ == "__main__":
    create_user_usecases_doc()
