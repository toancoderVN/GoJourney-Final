#!/usr/bin/env python3
"""Convert Agent Use Cases to DOCX format - Single table format"""

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def merge_cells(table, row_idx, col_start, col_end):
    """Merge cells in a row"""
    cell = table.cell(row_idx, col_start)
    for i in range(col_start + 1, col_end + 1):
        cell.merge(table.cell(row_idx, i))
    return cell

def add_usecase(doc, title, data):
    """Add a use case with single table format"""
    doc.add_heading(title, level=2)
    
    # Create table with 3 columns (Label | Actor | System)
    table = doc.add_table(rows=0, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    for item in data:
        row_type = item.get('type', 'single')
        
        if row_type == 'single':
            row = table.add_row()
            row.cells[0].text = item['label']
            row.cells[0].paragraphs[0].runs[0].bold = True
            
            merged = merge_cells(table, len(table.rows) - 1, 1, 2)
            merged.text = item['value']
            
        elif row_type == 'header_flow':
            row = table.add_row()
            row.cells[0].text = item['label']
            row.cells[0].paragraphs[0].runs[0].bold = True
            
            row.cells[1].text = 'Actor'
            row.cells[1].paragraphs[0].runs[0].bold = True
            
            row.cells[2].text = 'System'
            row.cells[2].paragraphs[0].runs[0].bold = True
            
        elif row_type == 'flow':
            row = table.add_row()
            row.cells[0].text = ''
            row.cells[1].text = item.get('actor', '')
            row.cells[2].text = item.get('system', '')
            
        elif row_type == 'exception_header':
            row = table.add_row()
            row.cells[0].text = item['label']
            row.cells[0].paragraphs[0].runs[0].bold = True
            
            merged = merge_cells(table, len(table.rows) - 1, 1, 2)
            merged.text = item['value']
            
        elif row_type == 'exception':
            row = table.add_row()
            row.cells[0].text = ''
            
            merged = merge_cells(table, len(table.rows) - 1, 1, 2)
            merged.text = item['value']
    
    doc.add_paragraph()

def create_agent_usecases_doc():
    doc = Document()
    
    # Title
    title = doc.add_heading('USE CASES - AI AGENT BOOKING', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('3.4.3. Module AI Booking Agent & Zalo Integration', level=1)
    
    # ==========================================
    # UC 3.4.3.1 - Tạo trip từ AI agent booking
    # ==========================================
    add_usecase(doc, '3.4.3.1. UseCase Tạo trip từ AI agent booking', [
        {'type': 'single', 'label': 'Usecase', 'value': 'Tạo trip từ AI agent booking'},
        {'type': 'single', 'label': 'Ngữ cảnh', 'value': 'Người dùng muốn nhờ AI tự động tìm kiếm và đặt phòng khách sạn theo yêu cầu'},
        {'type': 'single', 'label': 'Mô tả', 'value': 'Người dùng cung cấp thông tin chuyến đi (địa điểm, ngày đi, ngân sách), AI sẽ tìm kiếm và khởi tạo quy trình đặt phòng tự động'},
        {'type': 'single', 'label': 'Tác Nhân', 'value': 'Người dùng đã đăng nhập'},
        {'type': 'single', 'label': 'Sự kiện', 'value': 'Người dùng điền form "Bạn muốn đi đâu?" trên Dashboard'},
        {'type': 'single', 'label': 'Kết quả', 'value': 'Một Booking Session được khởi tạo và AI bắt đầu tìm kiếm khách sạn'},
        {'type': 'header_flow', 'label': 'Luồng sự kiện'},
        {'type': 'flow', 'actor': 'Người dùng nhập: Điểm đến, Ngày đi/về, Số người, Ngân sách dự kiến', 'system': 'Hệ thống hiển thị form nhập liệu thông minh'},
        {'type': 'flow', 'actor': 'Người dùng nhấn "Bắt đầu lập kế hoạch"', 'system': 'Hệ thống gọi API POST /chat/send-message với intent \'BOOKING_REQUEST\''},
        {'type': 'flow', 'actor': '', 'system': 'AI Service phân tích yêu cầu và tìm kiếm khách sạn phù hợp từ Database/RAG'},
        {'type': 'flow', 'actor': '', 'system': 'AI Service tạo một "Booking Session" mới với trạng thái "SEARCHING"'},
        {'type': 'flow', 'actor': '', 'system': 'Hệ thống chuyển hướng người dùng đến trang "Booking Mission" để theo dõi tiến trình'},
        {'type': 'exception_header', 'label': 'Ngoại lệ', 'value': 'Không tìm thấy địa điểm: AI thông báo "Tôi chưa có dữ liệu về địa điểm này, bạn có thể thử địa điểm khác không?"'},
        {'type': 'exception', 'value': 'Lỗi kết nối AI: Hệ thống báo "AI Service đang bận, vui lòng thử lại sau"'},
    ])
    
    # ==========================================
    # UC 3.4.3.2 - Đàm phán với khách sạn qua Zalo
    # ==========================================
    add_usecase(doc, '3.4.3.2. UseCase Đàm phán với khách sạn qua Zalo (tự động)', [
        {'type': 'single', 'label': 'Usecase', 'value': 'Đàm phán với khách sạn qua Zalo (tự động)'},
        {'type': 'single', 'label': 'Ngữ cảnh', 'value': 'AI Agent thay mặt người dùng liên hệ trực tiếp với khách sạn qua Zalo để hỏi giá và tình trạng phòng'},
        {'type': 'single', 'label': 'Mô tả', 'value': 'Sau khi người dùng chọn khách sạn, AI sẽ gửi tin nhắn Zalo tới khách sạn, xử lý phản hồi và báo cáo lại cho người dùng'},
        {'type': 'single', 'label': 'Tác Nhân', 'value': 'System (AI Agent)'},
        {'type': 'single', 'label': 'Sự kiện', 'value': 'Người dùng xác nhận chọn một khách sạn từ danh sách gợi ý của AI'},
        {'type': 'single', 'label': 'Kết quả', 'value': 'AI nhận được báo giá và tình trạng phòng từ khách sạn'},
        {'type': 'header_flow', 'label': 'Luồng sự kiện'},
        {'type': 'flow', 'actor': 'Người dùng nhấn "Chọn khách sạn này"', 'system': 'AI cập nhật trạng thái session thành "CONTACTING_HOTEL"'},
        {'type': 'flow', 'actor': '', 'system': 'AI Service gửi tin nhắn Zalo tới số điện thoại của khách sạn: "Xin chào, tôi muốn đặt phòng cho [Tên KH]..."'},
        {'type': 'flow', 'actor': '(Khách sạn) Nhắn tin phản hồi qua Zalo', 'system': 'Zalo Webhook nhận tin nhắn phản hồi và đẩy về AI Service'},
        {'type': 'flow', 'actor': '', 'system': 'AI phân tích nội dung tin nhắn (giá, phòng trống)'},
        {'type': 'flow', 'actor': '', 'system': 'AI cập nhật trạng thái thành "NEGOTIATING"'},
        {'type': 'flow', 'actor': '', 'system': 'AI thông báo cho người dùng trên Web App: "Khách sạn X đã phản hồi: Giá 1.2tr/đêm, còn phòng"'},
        {'type': 'exception_header', 'label': 'Ngoại lệ', 'value': 'Khách sạn không phản hồi: Sau 30 phút, AI gửi thông báo "Khách sạn chưa phản hồi, bạn có muốn thử khách sạn khác không?"'},
        {'type': 'exception', 'value': 'Lỗi gửi tin Zalo: Hệ thống log lỗi và báo cho Admin kiểm tra kết nối Zalo OA'},
    ])
    
    # ==========================================
    # UC 3.4.3.3 - Xác nhận thanh toán booking
    # ==========================================
    add_usecase(doc, '3.4.3.3. UseCase Xác nhận thanh toán booking', [
        {'type': 'single', 'label': 'Usecase', 'value': 'Xác nhận thanh toán booking'},
        {'type': 'single', 'label': 'Ngữ cảnh', 'value': 'Người dùng đồng ý với mức giá và điều kiện mà khách sạn đưa ra'},
        {'type': 'single', 'label': 'Mô tả', 'value': 'Người dùng xác nhận thanh toán để AI hoàn tất thủ tục đặt phòng với khách sạn'},
        {'type': 'single', 'label': 'Tác Nhân', 'value': 'Người dùng đã đăng nhập'},
        {'type': 'single', 'label': 'Sự kiện', 'value': 'Người dùng nhận được thông báo giá và nhấn "Xác nhận thanh toán"'},
        {'type': 'single', 'label': 'Kết quả', 'value': 'Booking được xác nhận, Trip được tạo chính thức trong hệ thống'},
        {'type': 'header_flow', 'label': 'Luồng sự kiện'},
        {'type': 'flow', 'actor': 'Người dùng xem chi tiết báo giá và nhấn "Thanh toán & Đặt phòng"', 'system': 'Hệ thống gọi API xác nhận booking'},
        {'type': 'flow', 'actor': '', 'system': 'AI gửi tin nhắn chốt đơn tới Zalo khách sạn: "Khách đã chốt, vui lòng giữ phòng."'},
        {'type': 'flow', 'actor': '', 'system': 'AI Service gọi Trip Service để tạo Trip mới và Booking record trong database'},
        {'type': 'flow', 'actor': '', 'system': 'Hệ thống cập nhật trạng thái session thành "CONFIRMED"'},
        {'type': 'flow', 'actor': '', 'system': 'Hệ thống hiển thị thông báo "Chúc mừng! Chuyến đi của bạn đã được lên lịch"'},
        {'type': 'exception_header', 'label': 'Ngoại lệ', 'value': 'Thanh toán thất bại: Hệ thống báo lỗi và giữ trạng thái "WAITING_PAYMENT"'},
        {'type': 'exception', 'value': 'Khách sạn hết phòng phút chót: AI thông báo và đề xuất khách sạn thay thế'},
    ])
    
    # ==========================================
    # UC 3.4.3.4 - Hủy booking request
    # ==========================================
    add_usecase(doc, '3.4.3.4. UseCase Hủy booking request', [
        {'type': 'single', 'label': 'Usecase', 'value': 'Hủy booking request'},
        {'type': 'single', 'label': 'Ngữ cảnh', 'value': 'Người dùng muốn hủy yêu cầu đặt phòng khi AI đang xử lý'},
        {'type': 'single', 'label': 'Mô tả', 'value': 'Người dùng dừng quy trình đặt phòng, AI sẽ dừng liên lạc với khách sạn'},
        {'type': 'single', 'label': 'Tác Nhân', 'value': 'Người dùng đã đăng nhập'},
        {'type': 'single', 'label': 'Sự kiện', 'value': 'Người dùng nhấn nút "Hủy yêu cầu" trên trang Booking Mission'},
        {'type': 'single', 'label': 'Kết quả', 'value': 'Booking session bị hủy, trạng thái chuyển về CANCELLED'},
        {'type': 'header_flow', 'label': 'Luồng sự kiện'},
        {'type': 'flow', 'actor': 'Người dùng nhấn nút "Hủy yêu cầu"', 'system': 'Hệ thống hiển thị popup xác nhận "Bạn có chắc muốn hủy?"'},
        {'type': 'flow', 'actor': 'Người dùng chọn "Đồng ý"', 'system': 'Hệ thống gọi API hủy booking session'},
        {'type': 'flow', 'actor': '', 'system': 'AI gửi tin nhắn xin lỗi tới khách sạn (nếu đã liên hệ): "Xin lỗi, khách đã thay đổi kế hoạch."'},
        {'type': 'flow', 'actor': '', 'system': 'Hệ thống cập nhật trạng thái session thành "CANCELLED"'},
        {'type': 'flow', 'actor': '', 'system': 'Hệ thống chuyển hướng người dùng về Dashboard'},
        {'type': 'exception_header', 'label': 'Ngoại lệ', 'value': 'Không thể hủy khi đã thanh toán: Hệ thống thông báo "Booking đã hoàn tất, vui lòng liên hệ CSKH để xử lý hoàn tiền"'},
    ])
    
    # Save document
    output_path = '/home/navin/Working/Booking_Agents/travel-agent-mvp/docs/usecases/agent_usecases.docx'
    doc.save(output_path)
    print(f"✅ Document saved to: {output_path}")

if __name__ == "__main__":
    create_agent_usecases_doc()
