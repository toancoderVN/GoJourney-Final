from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_trip_mgmt_doc():
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)

    md_path = 'docs/usecases/analysis_trip_management.md'
    if not os.path.exists(md_path):
        print(f"File {md_path} not found")
        return

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('# '):
            p = doc.add_paragraph(line.replace('# ', '').upper())
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0]
            run.bold = True
            run.font.size = Pt(16)
            doc.add_paragraph('')
            
        elif line.startswith('## '):
            p = doc.add_paragraph(line.replace('## ', ''))
            run = p.runs[0]
            run.bold = True
            run.font.size = Pt(14)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)

        elif line.startswith('### '):
            p = doc.add_paragraph(line.replace('### ', ''))
            run = p.runs[0]
            run.bold = True
            run.italic = True
            run.font.size = Pt(13)
            p.paragraph_format.space_before = Pt(6)
            
        elif line.startswith('*   ') or line.startswith('-   '):
            p = doc.add_paragraph(line.replace('*   ', '').replace('-   ', ''), style='List Bullet')
            
        elif line.startswith('*(Hình'):
            p = doc.add_paragraph(line.replace('*', ''))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.italic = True
            
        else:
            parts = line.split('**')
            p = doc.add_paragraph()
            for i, part in enumerate(parts):
                run = p.add_run(part)
                if i % 2 == 1:
                    run.bold = True
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    output_path = 'docs/usecases/analysis_trip_management.docx'
    doc.save(output_path)
    print(f"Created {output_path}")

if __name__ == "__main__":
    create_trip_mgmt_doc()
