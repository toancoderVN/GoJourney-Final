#!/usr/bin/env python3
"""Convert AI Research Use Cases to DOCX format - Single table format"""

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def merge_cells(table, row_idx, col_start, col_end):
    """Merge cells in a row"""
    cell = table.cell(row_idx, col_start)
    for i in range(col_start + 1, col_end + 1):
        cell.merge(table.cell(row_idx, i))
    return cell

def add_usecase(doc, title, data):
    """Add a use case with single table format"""
    doc.add_heading(title, level=2)
    
    table = doc.add_table(rows=0, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    for item in data:
        row_type = item.get('type', 'single')
        
        if row_type == 'single':
            row = table.add_row()
            row.cells[0].text = item['label']
            row.cells[0].paragraphs[0].runs[0].bold = True
            
            merged = merge_cells(table, len(table.rows) - 1, 1, 2)
            merged.text = item['value']
            
        elif row_type == 'header_flow':
            row = table.add_row()
            row.cells[0].text = item['label']
            row.cells[0].paragraphs[0].runs[0].bold = True
            
            row.cells[1].text = 'Actor'
            row.cells[1].paragraphs[0].runs[0].bold = True
            
            row.cells[2].text = 'System'
            row.cells[2].paragraphs[0].runs[0].bold = True
            
        elif row_type == 'flow':
            row = table.add_row()
            row.cells[0].text = ''
            row.cells[1].text = item.get('actor', '')
            row.cells[2].text = item.get('system', '')
            
        elif row_type == 'exception_header':
            row = table.add_row()
            row.cells[0].text = item['label']
            row.cells[0].paragraphs[0].runs[0].bold = True
            
            merged = merge_cells(table, len(table.rows) - 1, 1, 2)
            merged.text = item['value']
            
        elif row_type == 'exception':
            row = table.add_row()
            row.cells[0].text = ''
            
            merged = merge_cells(table, len(table.rows) - 1, 1, 2)
            merged.text = item['value']
    
    doc.add_paragraph()

def create_research_usecases_doc():
    doc = Document()
    
    # Title
    title = doc.add_heading('USE CASES - AI RESEARCH', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('3.4.5. Module AI Chat & Research', level=1)
    
    # ==========================================
    # UC 3.4.5.1 - Chat với AI assistant
    # ==========================================
    add_usecase(doc, '3.4.5.1. UseCase Chat với AI assistant', [
        {'type': 'single', 'label': 'Usecase', 'value': 'Chat với AI assistant'},
        {'type': 'single', 'label': 'Ngữ cảnh', 'value': 'Người dùng muốn trò chuyện với AI để hỏi đáp thông tin chung, lên kế hoạch du lịch hoặc giải trí'},
        {'type': 'single', 'label': 'Mô tả', 'value': 'Người dùng gửi tin nhắn văn bản tới AI, AI trả lời dựa trên kiến thức có sẵn và ngữ cảnh chuyến đi'},
        {'type': 'single', 'label': 'Tác Nhân', 'value': 'Người dùng đã đăng nhập'},
        {'type': 'single', 'label': 'Sự kiện', 'value': 'Người dùng nhập tin nhắn vào khung chat trên Dashboard hoặc trong Trip Detail'},
        {'type': 'single', 'label': 'Kết quả', 'value': 'AI trả lời câu hỏi và gợi ý các hành động tiếp theo'},
        {'type': 'header_flow', 'label': 'Luồng sự kiện'},
        {'type': 'flow', 'actor': 'Người dùng nhập: "Gợi ý lịch trình 3 ngày ở Đà Lạt"', 'system': 'Hệ thống hiển thị tin nhắn của người dùng'},
        {'type': 'flow', 'actor': 'Người dùng nhấn Gửi', 'system': 'Hệ thống gửi tin nhắn tới AI Service'},
        {'type': 'flow', 'actor': '', 'system': 'AI Service phân loại intent (General Chat)'},
        {'type': 'flow', 'actor': '', 'system': 'AI gọi LLM (GPT-4o) để generate câu trả lời'},
        {'type': 'flow', 'actor': '', 'system': 'Hệ thống hiển thị câu trả lời của AI kèm theo các gợi ý (Quick Actions)'},
        {'type': 'exception_header', 'label': 'Ngoại lệ', 'value': 'Lỗi kết nối LLM: Hệ thống báo "AI đang bận, vui lòng thử lại sau"'},
        {'type': 'exception', 'value': 'Nội dung vi phạm: AI từ chối trả lời nếu câu hỏi vi phạm chính sách an toàn'},
    ])
    
    # ==========================================
    # UC 3.4.5.2 - Nghiên cứu sâu về điểm đến
    # ==========================================
    add_usecase(doc, '3.4.5.2. UseCase Nghiên cứu sâu về điểm đến (Deep Research)', [
        {'type': 'single', 'label': 'Usecase', 'value': 'Nghiên cứu sâu về điểm đến (Deep Research)'},
        {'type': 'single', 'label': 'Ngữ cảnh', 'value': 'Người dùng cần thông tin chi tiết, chuyên sâu về một địa điểm du lịch (văn hóa, ẩm thực, lịch sử, tips du lịch)'},
        {'type': 'single', 'label': 'Mô tả', 'value': 'AI thực hiện quy trình nghiên cứu đa bước: thu thập thông tin từ nhiều nguồn, tổng hợp và phân tích để đưa ra báo cáo chi tiết'},
        {'type': 'single', 'label': 'Tác Nhân', 'value': 'Người dùng'},
        {'type': 'single', 'label': 'Sự kiện', 'value': 'Người dùng chọn chế độ "Deep Research" và nhập địa điểm/chủ đề quan tâm'},
        {'type': 'single', 'label': 'Kết quả', 'value': 'Một báo cáo chi tiết về điểm đến được hiển thị'},
        {'type': 'header_flow', 'label': 'Luồng sự kiện'},
        {'type': 'flow', 'actor': 'Người dùng chọn "Deep Research" cho địa điểm "Hang Sơn Đoòng"', 'system': 'Hệ thống khởi tạo Research Session'},
        {'type': 'flow', 'actor': '', 'system': 'AI xác định các khía cạnh cần nghiên cứu (Lịch sử, Cách đi, Chi phí, Lưu ý an toàn)'},
        {'type': 'flow', 'actor': '', 'system': 'AI thực hiện nhiều lượt tìm kiếm và tổng hợp thông tin từ Knowledge Base/Web'},
        {'type': 'flow', 'actor': '', 'system': 'AI tổng hợp thông tin thành một bài viết chi tiết, có cấu trúc'},
        {'type': 'flow', 'actor': '', 'system': 'Hệ thống hiển thị kết quả nghiên cứu dưới dạng bài viết rich-text'},
        {'type': 'exception_header', 'label': 'Ngoại lệ', 'value': 'Thiếu dữ liệu: AI thông báo "Không đủ dữ liệu tin cậy để nghiên cứu sâu về địa điểm này"'},
        {'type': 'exception', 'value': 'Quá tải request: Hệ thống báo "Vui lòng chờ giây lát trước khi thực hiện nghiên cứu mới"'},
    ])
    
    # ==========================================
    # UC 3.4.5.3 - Tìm kiếm thông tin du lịch trên web
    # ==========================================
    add_usecase(doc, '3.4.5.3. UseCase Tìm kiếm thông tin du lịch trên web', [
        {'type': 'single', 'label': 'Usecase', 'value': 'Tìm kiếm thông tin du lịch trên web'},
        {'type': 'single', 'label': 'Ngữ cảnh', 'value': 'Người dùng cần thông tin cập nhật mới nhất (thời tiết, giá vé hiện tại, sự kiện sắp diễn ra) mà AI không có sẵn'},
        {'type': 'single', 'label': 'Mô tả', 'value': 'AI sử dụng công cụ tìm kiếm web để lấy thông tin real-time và trả lời câu hỏi của người dùng'},
        {'type': 'single', 'label': 'Tác Nhân', 'value': 'Người dùng'},
        {'type': 'single', 'label': 'Sự kiện', 'value': 'Người dùng hỏi về thông tin mang tính thời sự (VD: "Thời tiết Đà Nẵng tuần tới thế nào?")'},
        {'type': 'single', 'label': 'Kết quả', 'value': 'Thông tin mới nhất được cập nhật và trả lời cho người dùng'},
        {'type': 'header_flow', 'label': 'Luồng sự kiện'},
        {'type': 'flow', 'actor': 'Người dùng hỏi: "Giá vé máy bay Hà Nội - HCM hôm nay bao nhiêu?"', 'system': 'AI phân tích intent và nhận thấy cần "Web Search"'},
        {'type': 'flow', 'actor': '', 'system': 'AI thực hiện query tới Search Engine (Tavily/Google API)'},
        {'type': 'flow', 'actor': '', 'system': 'AI đọc và trích xuất thông tin từ các kết quả tìm kiếm hàng đầu'},
        {'type': 'flow', 'actor': '', 'system': 'AI tổng hợp câu trả lời dựa trên dữ liệu vừa tìm được, kèm trích dẫn nguồn'},
        {'type': 'flow', 'actor': '', 'system': 'Hệ thống hiển thị câu trả lời và link nguồn tham khảo'},
        {'type': 'exception_header', 'label': 'Ngoại lệ', 'value': 'Không tìm thấy kết quả: AI báo "Không tìm thấy thông tin phù hợp trên web"'},
        {'type': 'exception', 'value': 'Tìm kiếm timeout: Hệ thống trả về câu trả lời dựa trên kiến thức có sẵn và cảnh báo "Dữ liệu có thể chưa cập nhật"'},
    ])
    
    # Save document
    output_path = '/home/navin/Working/Booking_Agents/travel-agent-mvp/docs/usecases/ai_research_usecases.docx'
    doc.save(output_path)
    print(f"✅ Document saved to: {output_path}")

if __name__ == "__main__":
    create_research_usecases_doc()
