#!/usr/bin/env python3
"""Convert Companion Management Use Cases to DOCX format - Single table format"""

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def merge_cells(table, row_idx, col_start, col_end):
    """Merge cells in a row"""
    cell = table.cell(row_idx, col_start)
    for i in range(col_start + 1, col_end + 1):
        cell.merge(table.cell(row_idx, i))
    return cell

def add_usecase(doc, title, data):
    """Add a use case with single table format"""
    doc.add_heading(title, level=2)
    
    table = doc.add_table(rows=0, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    for item in data:
        row_type = item.get('type', 'single')
        
        if row_type == 'single':
            row = table.add_row()
            row.cells[0].text = item['label']
            row.cells[0].paragraphs[0].runs[0].bold = True
            
            merged = merge_cells(table, len(table.rows) - 1, 1, 2)
            merged.text = item['value']
            
        elif row_type == 'header_flow':
            row = table.add_row()
            row.cells[0].text = item['label']
            row.cells[0].paragraphs[0].runs[0].bold = True
            
            row.cells[1].text = 'Actor'
            row.cells[1].paragraphs[0].runs[0].bold = True
            
            row.cells[2].text = 'System'
            row.cells[2].paragraphs[0].runs[0].bold = True
            
        elif row_type == 'flow':
            row = table.add_row()
            row.cells[0].text = ''
            row.cells[1].text = item.get('actor', '')
            row.cells[2].text = item.get('system', '')
            
        elif row_type == 'exception_header':
            row = table.add_row()
            row.cells[0].text = item['label']
            row.cells[0].paragraphs[0].runs[0].bold = True
            
            merged = merge_cells(table, len(table.rows) - 1, 1, 2)
            merged.text = item['value']
            
        elif row_type == 'exception':
            row = table.add_row()
            row.cells[0].text = ''
            
            merged = merge_cells(table, len(table.rows) - 1, 1, 2)
            merged.text = item['value']
    
    doc.add_paragraph()

def create_companion_usecases_doc():
    doc = Document()
    
    # Title
    title = doc.add_heading('USE CASES - COMPANION MANAGEMENT', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('3.4.7. Module Quản Lý Bạn Đồng Hành (Chi tiết)', level=1)
    
    # ==========================================
    # UC 3.4.7.1 - Xem danh sách người đồng hành
    # ==========================================
    add_usecase(doc, '3.4.7.1. UseCase Xem danh sách người đồng hành', [
        {'type': 'single', 'label': 'Usecase', 'value': 'Xem danh sách người đồng hành'},
        {'type': 'single', 'label': 'Ngữ cảnh', 'value': 'Người dùng muốn quản lý danh sách bạn bè, người thân đã kết nối để mời vào các chuyến đi'},
        {'type': 'single', 'label': 'Mô tả', 'value': 'Hiển thị danh sách chi tiết các companion, bao gồm thông tin cá nhân cơ bản và lịch sử chuyến đi chung'},
        {'type': 'single', 'label': 'Tác Nhân', 'value': 'Người dùng đã đăng nhập'},
        {'type': 'single', 'label': 'Sự kiện', 'value': 'Người dùng truy cập trang "Companions" từ Dashboard'},
        {'type': 'single', 'label': 'Kết quả', 'value': 'Danh sách companion hiển thị đầy đủ'},
        {'type': 'header_flow', 'label': 'Luồng sự kiện'},
        {'type': 'flow', 'actor': 'Người dùng chọn menu "Companions"', 'system': 'Hệ thống gọi API GET /travel-companions'},
        {'type': 'flow', 'actor': '', 'system': 'Hệ thống truy vấn database lấy danh sách bạn bè đã chấp nhận kết nối'},
        {'type': 'flow', 'actor': '', 'system': 'Hệ thống hiển thị danh sách dạng thẻ (Card view) hoặc danh sách (List view)'},
        {'type': 'flow', 'actor': '', 'system': 'Hiển thị các nút thao tác nhanh: "Xóa", "Mời chuyến đi"'},
        {'type': 'exception_header', 'label': 'Ngoại lệ', 'value': 'Lỗi tải dữ liệu: Hệ thống hiển thị "Error loading companions" và nút Retry'},
    ])
    
    # ==========================================
    # UC 3.4.7.2 - Kết nối với người dùng khác qua User ID
    # ==========================================
    add_usecase(doc, '3.4.7.2. UseCase Kết nối với người dùng khác qua User ID', [
        {'type': 'single', 'label': 'Usecase', 'value': 'Kết nối với người dùng khác qua User ID'},
        {'type': 'single', 'label': 'Ngữ cảnh', 'value': 'Người dùng biết User ID (Internal ID) của người khác và muốn gửi lời mời kết bạn trực tiếp'},
        {'type': 'single', 'label': 'Mô tả', 'value': 'Gửi lời mời kết bạn thông qua việc nhập chính xác User ID của đối phương'},
        {'type': 'single', 'label': 'Tác Nhân', 'value': 'Người dùng đã đăng nhập'},
        {'type': 'single', 'label': 'Sự kiện', 'value': 'Người dùng chọn "Thêm bạn bằng ID"'},
        {'type': 'single', 'label': 'Kết quả', 'value': 'Lời mời kết bạn được gửi đi thành công'},
        {'type': 'header_flow', 'label': 'Luồng sự kiện'},
        {'type': 'flow', 'actor': 'Người dùng chọn "Add Companion" -> "By User ID"', 'system': 'Hệ thống hiển thị modal nhập User ID'},
        {'type': 'flow', 'actor': 'Người dùng nhập ID (UUID) và nhấn "Gửi lời mời"', 'system': 'Hệ thống gọi API POST /travel-companions/invite'},
        {'type': 'flow', 'actor': '', 'system': 'Hệ thống kiểm tra User ID có tồn tại'},
        {'type': 'flow', 'actor': '', 'system': 'Hệ thống kiểm tra xem đã là bạn bè chưa'},
        {'type': 'flow', 'actor': '', 'system': 'Hệ thống tạo Invitation record và gửi notification'},
        {'type': 'flow', 'actor': '', 'system': 'Thông báo "Invitation sent successfully"'},
        {'type': 'exception_header', 'label': 'Ngoại lệ', 'value': 'User ID không đúng định dạng: Hệ thống báo "Invalid User ID format"'},
        {'type': 'exception', 'value': 'Không tìm thấy User: Hệ thống báo "User not found"'},
    ])
    
    # ==========================================
    # UC 3.4.7.3 - Xem danh sách lời mời chờ xử lý (Pending)
    # ==========================================
    add_usecase(doc, '3.4.7.3. UseCase Xem danh sách lời mời chờ xử lý', [
        {'type': 'single', 'label': 'Usecase', 'value': 'Xem danh sách lời mời chờ xử lý'},
        {'type': 'single', 'label': 'Ngữ cảnh', 'value': 'Người dùng kiểm tra các lời mời kết bạn (Inbound) và các lời mời mình đã gửi đi (Outbound)'},
        {'type': 'single', 'label': 'Mô tả', 'value': 'Quản lý các lời mời chưa được chấp nhận, cho phép chấp nhận/từ chối hoặc hủy lời mời đã gửi'},
        {'type': 'single', 'label': 'Tác Nhân', 'value': 'Người dùng đã đăng nhập'},
        {'type': 'single', 'label': 'Sự kiện', 'value': 'Người dùng chọn tab "Pending Invitations"'},
        {'type': 'single', 'label': 'Kết quả', 'value': 'Danh sách lời mời Inbound/Outbound hiển thị'},
        {'type': 'header_flow', 'label': 'Luồng sự kiện'},
        {'type': 'flow', 'actor': 'Người dùng chọn tab "Pending"', 'system': 'Hệ thống gọi API GET /travel-companions/invitations'},
        {'type': 'flow', 'actor': '', 'system': 'Hệ thống phân loại lời mời: "Received" (Đã nhận) và "Sent" (Đã gửi)'},
        {'type': 'flow', 'actor': '', 'system': 'Hiển thị danh sách với các nút hành động tương ứng (Accept/Decline cho Received, Cancel cho Sent)'},
        {'type': 'exception_header', 'label': 'Ngoại lệ', 'value': 'Danh sách trống: Hiển thị "No pending invitations"'},
    ])
    
    # ==========================================
    # UC 3.4.7.4 - Xóa người đồng hành
    # ==========================================
    add_usecase(doc, '3.4.7.4. UseCase Xóa người đồng hành', [
        {'type': 'single', 'label': 'Usecase', 'value': 'Xóa người đồng hành'},
        {'type': 'single', 'label': 'Ngữ cảnh', 'value': 'Người dùng muốn hủy kết nối bạn bè với một người khác'},
        {'type': 'single', 'label': 'Mô tả', 'value': 'Xóa mối quan hệ companion khỏi hệ thống, người kia sẽ không còn thấy người dùng trong danh sách bạn bè'},
        {'type': 'single', 'label': 'Tác Nhân', 'value': 'Người dùng đã đăng nhập'},
        {'type': 'single', 'label': 'Sự kiện', 'value': 'Người dùng nhấn nút "Xóa" (hoặc icon thùng rác) trên thẻ thông tin companion'},
        {'type': 'single', 'label': 'Kết quả', 'value': 'Mối quan hệ bạn bè bị xóa vĩnh viễn'},
        {'type': 'header_flow', 'label': 'Luồng sự kiện'},
        {'type': 'flow', 'actor': 'Người dùng nhấn nút "Remove Companion"', 'system': 'Hệ thống hiển thị dialog xác nhận "Are you sure?"'},
        {'type': 'flow', 'actor': 'Người dùng chọn "Yes, remove"', 'system': 'Hệ thống gọi API DELETE /travel-companions/:id'},
        {'type': 'flow', 'actor': '', 'system': 'Hệ thống xóa record trong bảng travel_companions (cả 2 chiều)'},
        {'type': 'flow', 'actor': '', 'system': 'Hệ thống gửi thông báo cho người dùng kia (tùy chọn)'},
        {'type': 'flow', 'actor': '', 'system': 'Hệ thống cập nhật lại danh sách hiển thị (bỏ item vừa xóa)'},
        {'type': 'exception_header', 'label': 'Ngoại lệ', 'value': 'Lỗi server: Hệ thống báo "Could not remove companion"'},
    ])
    
    # Save document
    output_path = '/home/navin/Working/Booking_Agents/travel-agent-mvp/docs/usecases/companion_management_usecases.docx'
    doc.save(output_path)
    print(f"✅ Document saved to: {output_path}")

if __name__ == "__main__":
    create_companion_usecases_doc()
