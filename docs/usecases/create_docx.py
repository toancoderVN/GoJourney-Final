#!/usr/bin/env python3
"""Convert Auth Use Cases to DOCX format - Single table format"""

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    """Set cell background color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def merge_cells(table, row_idx, col_start, col_end):
    """Merge cells in a row"""
    cell = table.cell(row_idx, col_start)
    for i in range(col_start + 1, col_end + 1):
        cell.merge(table.cell(row_idx, i))
    return cell

def add_usecase(doc, title, data):
    """Add a use case with single table format"""
    doc.add_heading(title, level=2)
    
    # Create table with 3 columns (for Actor/System in Luồng sự kiện)
    table = doc.add_table(rows=0, cols=3)
    table.style = 'Table Grid'
    
    for item in data:
        row_type = item.get('type', 'single')
        
        if row_type == 'single':
            # Single row with label and value (merged cols 1-2)
            row = table.add_row()
            row.cells[0].text = item['label']
            row.cells[0].paragraphs[0].runs[0].bold = True
            merged = merge_cells(table, len(table.rows) - 1, 1, 2)
            merged.text = item['value']
            
        elif row_type == 'header_flow':
            # Luồng sự kiện header row
            row = table.add_row()
            row.cells[0].text = item['label']
            row.cells[0].paragraphs[0].runs[0].bold = True
            row.cells[1].text = 'Actor'
            row.cells[1].paragraphs[0].runs[0].bold = True
            row.cells[2].text = 'System'
            row.cells[2].paragraphs[0].runs[0].bold = True
            
        elif row_type == 'flow':
            # Flow row with Actor and System
            row = table.add_row()
            row.cells[0].text = ''
            row.cells[1].text = item.get('actor', '')
            row.cells[2].text = item.get('system', '')
            
        elif row_type == 'exception_header':
            # Ngoại lệ header
            row = table.add_row()
            row.cells[0].text = item['label']
            row.cells[0].paragraphs[0].runs[0].bold = True
            merged = merge_cells(table, len(table.rows) - 1, 1, 2)
            merged.text = item['value']
            
        elif row_type == 'exception':
            # Exception row
            row = table.add_row()
            row.cells[0].text = ''
            merged = merge_cells(table, len(table.rows) - 1, 1, 2)
            merged.text = item['value']
    
    doc.add_paragraph()

def create_auth_usecases_doc():
    doc = Document()
    
    # Title
    title = doc.add_heading('USE CASES - AUTH SERVICE', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('3.4.1. Module Xác Thực (Authentication)', level=1)
    
    # ==========================================
    # UC 3.4.1.1 - Đăng ký tài khoản mới
    # ==========================================
    add_usecase(doc, '3.4.1.1. UseCase Đăng ký tài khoản mới', [
        {'type': 'single', 'label': 'Usecase', 'value': 'Đăng ký tài khoản mới'},
        {'type': 'single', 'label': 'Ngữ cảnh', 'value': 'Quản lý đăng ký tài khoản cho người dùng hệ thống Travel Agent'},
        {'type': 'single', 'label': 'Mô tả', 'value': 'Người dùng tạo tài khoản mới để sử dụng các chức năng đặt phòng, quản lý chuyến đi và chat với AI Agent'},
        {'type': 'single', 'label': 'Tác Nhân', 'value': 'Khách (Guest)'},
        {'type': 'single', 'label': 'Sự kiện', 'value': 'Người dùng truy cập hệ thống và chọn đăng ký tài khoản mới'},
        {'type': 'single', 'label': 'Kết quả', 'value': 'Tài khoản được tạo thành công, người dùng được đăng nhập tự động và chuyển đến Dashboard'},
        {'type': 'header_flow', 'label': 'Luồng sự kiện'},
        {'type': 'flow', 'actor': 'Người dùng truy cập trang đăng ký', 'system': 'Hệ thống hiển thị form đăng ký với các trường: Email, Mật khẩu, Xác nhận mật khẩu, Họ, Tên'},
        {'type': 'flow', 'actor': 'Người dùng nhập thông tin đăng ký', 'system': ''},
        {'type': 'flow', 'actor': 'Người dùng nhấn nút "Đăng ký"', 'system': 'Hệ thống kiểm tra tính hợp lệ của thông tin'},
        {'type': 'flow', 'actor': '', 'system': 'Hệ thống kiểm tra email chưa tồn tại trong hệ thống'},
        {'type': 'flow', 'actor': '', 'system': 'Hệ thống tạo tài khoản mới và mã hóa mật khẩu'},
        {'type': 'flow', 'actor': '', 'system': 'Hệ thống tạo JWT access token và refresh token'},
        {'type': 'flow', 'actor': '', 'system': 'Hệ thống chuyển hướng đến Dashboard'},
        {'type': 'exception_header', 'label': 'Ngoại lệ', 'value': 'Email đã tồn tại: Hệ thống báo lỗi "Email đã được sử dụng" và yêu cầu nhập email khác'},
        {'type': 'exception', 'value': 'Mật khẩu không khớp: Hệ thống báo lỗi "Mật khẩu xác nhận không khớp"'},
        {'type': 'exception', 'value': 'Email không đúng định dạng: Hệ thống báo lỗi "Email không hợp lệ"'},
    ])
    
    # ==========================================
    # UC 3.4.1.2 - Đăng nhập
    # ==========================================
    add_usecase(doc, '3.4.1.2. UseCase Đăng nhập', [
        {'type': 'single', 'label': 'Usecase', 'value': 'Đăng nhập'},
        {'type': 'single', 'label': 'Ngữ cảnh', 'value': 'Quản lý xác thực người dùng cho hệ thống Travel Agent'},
        {'type': 'single', 'label': 'Mô tả', 'value': 'Đăng nhập để sử dụng các chức năng đặt phòng, quản lý chuyến đi, chat với AI Agent'},
        {'type': 'single', 'label': 'Tác Nhân', 'value': 'Người dùng đã có tài khoản'},
        {'type': 'single', 'label': 'Sự kiện', 'value': 'Người dùng truy cập hệ thống và chọn đăng nhập'},
        {'type': 'single', 'label': 'Kết quả', 'value': 'Đăng nhập thành công và hiển thị Dashboard với các chức năng của hệ thống'},
        {'type': 'header_flow', 'label': 'Luồng sự kiện'},
        {'type': 'flow', 'actor': 'Người dùng chọn mục "Đăng nhập"', 'system': 'Hệ thống hiển thị giao diện đăng nhập'},
        {'type': 'flow', 'actor': 'Người dùng nhập Email và Mật khẩu', 'system': ''},
        {'type': 'flow', 'actor': 'Người dùng nhấn nút "Đăng nhập"', 'system': 'Hệ thống kiểm tra thông tin và xác thực'},
        {'type': 'flow', 'actor': '', 'system': 'Nếu hợp lệ, hệ thống tạo JWT access token và refresh token'},
        {'type': 'flow', 'actor': '', 'system': 'Hệ thống lưu tokens vào localStorage'},
        {'type': 'flow', 'actor': '', 'system': 'Hệ thống hiển thị Dashboard và các chức năng'},
        {'type': 'exception_header', 'label': 'Ngoại lệ', 'value': 'Người dùng nhập sai tài khoản hoặc mật khẩu, hệ thống báo lỗi và yêu cầu nhập lại'},
    ])
    
    # ==========================================
    # UC 3.4.1.3 - Đăng xuất
    # ==========================================
    add_usecase(doc, '3.4.1.3. UseCase Đăng xuất', [
        {'type': 'single', 'label': 'Usecase', 'value': 'Đăng xuất'},
        {'type': 'single', 'label': 'Ngữ cảnh', 'value': 'Kết thúc phiên làm việc của người dùng trong hệ thống'},
        {'type': 'single', 'label': 'Mô tả', 'value': 'Người dùng đăng xuất khỏi hệ thống để bảo mật thông tin cá nhân'},
        {'type': 'single', 'label': 'Tác Nhân', 'value': 'Người dùng đã đăng nhập'},
        {'type': 'single', 'label': 'Sự kiện', 'value': 'Người dùng nhấn nút đăng xuất trên thanh navigation'},
        {'type': 'single', 'label': 'Kết quả', 'value': 'Phiên đăng nhập kết thúc, người dùng được chuyển về trang đăng nhập'},
        {'type': 'header_flow', 'label': 'Luồng sự kiện'},
        {'type': 'flow', 'actor': 'Người dùng nhấn vào avatar/menu người dùng', 'system': 'Hệ thống hiển thị dropdown menu'},
        {'type': 'flow', 'actor': 'Người dùng chọn "Đăng xuất"', 'system': 'Hệ thống xóa JWT tokens khỏi localStorage'},
        {'type': 'flow', 'actor': '', 'system': 'Hệ thống ngắt kết nối WebSocket'},
        {'type': 'flow', 'actor': '', 'system': 'Hệ thống chuyển hướng về trang đăng nhập'},
        {'type': 'exception_header', 'label': 'Ngoại lệ', 'value': 'Lỗi kết nối mạng: Hệ thống vẫn xóa tokens local và chuyển về trang đăng nhập'},
    ])
    
    # ==========================================
    # UC 3.4.1.4 - Xem thông tin profile
    # ==========================================
    add_usecase(doc, '3.4.1.4. UseCase Xem thông tin profile người dùng', [
        {'type': 'single', 'label': 'Usecase', 'value': 'Xem thông tin profile người dùng'},
        {'type': 'single', 'label': 'Ngữ cảnh', 'value': 'Quản lý thông tin cá nhân của người dùng trong hệ thống'},
        {'type': 'single', 'label': 'Mô tả', 'value': 'Người dùng xem và quản lý thông tin cá nhân của mình bao gồm họ tên, email, avatar'},
        {'type': 'single', 'label': 'Tác Nhân', 'value': 'Người dùng đã đăng nhập'},
        {'type': 'single', 'label': 'Sự kiện', 'value': 'Người dùng truy cập trang Profile hoặc nhấn vào avatar'},
        {'type': 'single', 'label': 'Kết quả', 'value': 'Hiển thị thông tin chi tiết của người dùng'},
        {'type': 'header_flow', 'label': 'Luồng sự kiện'},
        {'type': 'flow', 'actor': 'Người dùng nhấn vào menu "Hồ sơ" hoặc avatar', 'system': 'Hệ thống gọi API GET /auth/profile với JWT token'},
        {'type': 'flow', 'actor': '', 'system': 'Hệ thống lấy thông tin người dùng từ database'},
        {'type': 'flow', 'actor': '', 'system': 'Hệ thống hiển thị trang Profile với thông tin: Avatar, Họ và Tên, Email, Ngày tạo tài khoản'},
        {'type': 'exception_header', 'label': 'Ngoại lệ', 'value': 'Token hết hạn: Hệ thống tự động refresh token hoặc yêu cầu đăng nhập lại'},
    ])
    
    # Save document
    output_path = '/home/navin/Working/Booking_Agents/travel-agent-mvp/docs/usecases/auth_usecases.docx'
    doc.save(output_path)
    print(f"✅ Document saved to: {output_path}")

if __name__ == "__main__":
    create_auth_usecases_doc()
