#!/usr/bin/env python3
"""Convert Zalo Use Cases to DOCX format - Single table format"""

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def merge_cells(table, row_idx, col_start, col_end):
    """Merge cells in a row"""
    cell = table.cell(row_idx, col_start)
    for i in range(col_start + 1, col_end + 1):
        cell.merge(table.cell(row_idx, i))
    return cell

def add_usecase(doc, title, data):
    """Add a use case with single table format"""
    doc.add_heading(title, level=2)
    
    table = doc.add_table(rows=0, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    for item in data:
        row_type = item.get('type', 'single')
        
        if row_type == 'single':
            row = table.add_row()
            row.cells[0].text = item['label']
            row.cells[0].paragraphs[0].runs[0].bold = True
            
            merged = merge_cells(table, len(table.rows) - 1, 1, 2)
            merged.text = item['value']
            
        elif row_type == 'header_flow':
            row = table.add_row()
            row.cells[0].text = item['label']
            row.cells[0].paragraphs[0].runs[0].bold = True
            
            row.cells[1].text = 'Actor'
            row.cells[1].paragraphs[0].runs[0].bold = True
            
            row.cells[2].text = 'System'
            row.cells[2].paragraphs[0].runs[0].bold = True
            
        elif row_type == 'flow':
            row = table.add_row()
            row.cells[0].text = ''
            row.cells[1].text = item.get('actor', '')
            row.cells[2].text = item.get('system', '')
            
        elif row_type == 'exception_header':
            row = table.add_row()
            row.cells[0].text = item['label']
            row.cells[0].paragraphs[0].runs[0].bold = True
            
            merged = merge_cells(table, len(table.rows) - 1, 1, 2)
            merged.text = item['value']
            
        elif row_type == 'exception':
            row = table.add_row()
            row.cells[0].text = ''
            
            merged = merge_cells(table, len(table.rows) - 1, 1, 2)
            merged.text = item['value']
    
    doc.add_paragraph()

def create_zalo_usecases_doc():
    doc = Document()
    
    # Title
    title = doc.add_heading('USE CASES - ZALO INTEGRATION', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('3.4.4. Module Zalo Integration', level=1)
    
    # ==========================================
    # UC 3.4.4.1 - Kết nối tài khoản Zalo
    # ==========================================
    add_usecase(doc, '3.4.4.1. UseCase Kết nối tài khoản Zalo', [
        {'type': 'single', 'label': 'Usecase', 'value': 'Kết nối tài khoản Zalo'},
        {'type': 'single', 'label': 'Ngữ cảnh', 'value': 'Người dùng hoặc Admin kết nối tài khoản Zalo Cá nhân/OA vào hệ thống để AI có thể thay mặt nhắn tin'},
        {'type': 'single', 'label': 'Mô tả', 'value': 'Thiết lập kết nối giữa hệ thống Travel Agent và tài khoản Zalo thông qua quét mã QR hoặc cấu hình OA'},
        {'type': 'single', 'label': 'Tác Nhân', 'value': 'User/Admin'},
        {'type': 'single', 'label': 'Sự kiện', 'value': 'Người dùng truy cập trang "Zalo Settings" và chọn "Kết nối Zalo"'},
        {'type': 'single', 'label': 'Kết quả', 'value': 'Tài khoản Zalo được kết nối thành công, hệ thống bắt đầu lắng nghe tin nhắn'},
        {'type': 'header_flow', 'label': 'Luồng sự kiện'},
        {'type': 'flow', 'actor': 'Người dùng nhấn "Kết nối tài khoản Zalo"', 'system': 'Hệ thống hiển thị mã QR Code (từ thư viện zalo-chat-client hoặc OA Auth URL)'},
        {'type': 'flow', 'actor': 'Người dùng mở app Zalo và quét mã', 'system': 'Hệ thống chờ xác thực từ Zalo server'},
        {'type': 'flow', 'actor': '', 'system': 'Khi xác thực thành công, hệ thống lưu credentials (cookies/tokens) vào secure storage'},
        {'type': 'flow', 'actor': '', 'system': 'Hệ thống tự động khởi chạy "Zalo Listener" để lắng nghe tin nhắn'},
        {'type': 'flow', 'actor': '', 'system': 'Hệ thống thông báo "Kết nối thành công: [Tên Zalo]"'},
        {'type': 'exception_header', 'label': 'Ngoại lệ', 'value': 'Quét mã thất bại/timeout: Hệ thống báo "Hết thời gian chờ, vui lòng thử lại" và refresh QR'},
    ])
    
    # ==========================================
    # UC 3.4.4.2 - Xem danh sách Zalo conversations
    # ==========================================
    add_usecase(doc, '3.4.4.2. UseCase Xem danh sách Zalo conversations', [
        {'type': 'single', 'label': 'Usecase', 'value': 'Xem danh sách Zalo conversations'},
        {'type': 'single', 'label': 'Ngữ cảnh', 'value': 'Quản lý các cuộc hội thoại đang diễn ra giữa AI và các khách sạn'},
        {'type': 'single', 'label': 'Mô tả', 'value': 'Xem danh sách tất cả các cuộc trò chuyện trên Zalo được đồng bộ về Dashboard'},
        {'type': 'single', 'label': 'Tác Nhân', 'value': 'Admin/User'},
        {'type': 'single', 'label': 'Sự kiện', 'value': 'Người dùng truy cập tab "Zalo Conversations"'},
        {'type': 'single', 'label': 'Kết quả', 'value': 'Danh sách hội thoại hiển thị với tin nhắn mới nhất'},
        {'type': 'header_flow', 'label': 'Luồng sự kiện'},
        {'type': 'flow', 'actor': 'Người dùng chọn mục "Zalo Conversations"', 'system': 'Hệ thống gọi API GET /zalo/conversations'},
        {'type': 'flow', 'actor': '', 'system': 'Hệ thống truy vấn danh sách hội thoại từ database (đã sync từ Zalo)'},
        {'type': 'flow', 'actor': '', 'system': 'Hiển thị danh sách: Avatar đối tác, Tên, Tin nhắn cuối, Thời gian'},
        {'type': 'exception_header', 'label': 'Ngoại lệ', 'value': 'Chưa kết nối Zalo: Hệ thống nhắc "Vui lòng kết nối tài khoản Zalo trước"'},
        {'type': 'exception', 'value': 'Lỗi đồng bộ: Hệ thống hiển thị "Không thể tải tin nhắn mới nhất"'},
    ])
    
    # ==========================================
    # UC 3.4.4.3 - Gửi tin nhắn qua Zalo tới khách sạn
    # ==========================================
    add_usecase(doc, '3.4.4.3. UseCase Gửi tin nhắn qua Zalo tới khách sạn', [
        {'type': 'single', 'label': 'Usecase', 'value': 'Gửi tin nhắn qua Zalo tới khách sạn'},
        {'type': 'single', 'label': 'Ngữ cảnh', 'value': 'AI Agent hoặc người dùng gửi tin nhắn hỏi giá/đặt phòng tới khách sạn'},
        {'type': 'single', 'label': 'Mô tả', 'value': 'Hệ thống gửi tin nhắn văn bản thông qua Zalo API tới số điện thoại của khách sạn'},
        {'type': 'single', 'label': 'Tác Nhân', 'value': 'System (AI Agent) / User'},
        {'type': 'single', 'label': 'Sự kiện', 'value': 'AI cần liên hệ khách sạn HOẶC người dùng nhập tin nhắn thủ công'},
        {'type': 'single', 'label': 'Kết quả', 'value': 'Tin nhắn được gửi đi và hiển thị trạng thái "Đã gửi"'},
        {'type': 'header_flow', 'label': 'Luồng sự kiện'},
        {'type': 'flow', 'actor': 'Hệ thống kích hoạt gửi tin (theo kịch bản AI hoặc user input)', 'system': 'AI Service xác định số điện thoại người nhận'},
        {'type': 'flow', 'actor': '', 'system': 'Hệ thống gọi Zalo API để gửi tin nhắn'},
        {'type': 'flow', 'actor': '', 'system': 'Hệ thống lưu tin nhắn vào database với trạng thái "SENT"'},
        {'type': 'flow', 'actor': '', 'system': 'Hệ thống cập nhật giao diện chat realtime'},
        {'type': 'exception_header', 'label': 'Ngoại lệ', 'value': 'Số điện thoại chưa kích hoạt Zalo: Hệ thống báo lỗi "Số điện thoại này chưa đăng ký Zalo"'},
        {'type': 'exception', 'value': 'Tài khoản bị khóa/chặn: Hệ thống nhận lỗi từ Zalo API và thông báo "Không thể gửi tin tin nhắn (Bị chặn/Lỗi)"'},
    ])
    
    # ==========================================
    # UC 3.4.4.4 - Nhận tin nhắn phản hồi từ khách sạn
    # ==========================================
    add_usecase(doc, '3.4.4.4. UseCase Nhận tin nhắn phản hồi từ khách sạn', [
        {'type': 'single', 'label': 'Usecase', 'value': 'Nhận tin nhắn phản hồi từ khách sạn'},
        {'type': 'single', 'label': 'Ngữ cảnh', 'value': 'Khách sạn trả lời tin nhắn của AI (báo giá, xác nhận chỗ trống)'},
        {'type': 'single', 'label': 'Mô tả', 'value': 'Hệ thống lắng nghe tin nhắn đến từ Zalo, lưu trữ và phân tích nội dung (nếu là AI booking flow)'},
        {'type': 'single', 'label': 'Tác Nhân', 'value': 'Khách sạn (External)'},
        {'type': 'single', 'label': 'Sự kiện', 'value': 'Khách sạn gửi tin nhắn tới tài khoản Zalo của hệ thống'},
        {'type': 'single', 'label': 'Kết quả', 'value': 'Tin nhắn được lưu vào hệ thống và notify cho người dùng/AI'},
        {'type': 'header_flow', 'label': 'Luồng sự kiện'},
        {'type': 'flow', 'actor': 'Khách sạn gửi tin nhắn "Giá phòng là 1tr/đêm"', 'system': 'Zalo Webhook/Listener nhận sự kiện tin nhắn mới'},
        {'type': 'flow', 'actor': '', 'system': 'Hệ thống lưu nội dung tin nhắn vào database'},
        {'type': 'flow', 'actor': '', 'system': 'Hệ thống kiểm tra xem tin nhắn này có thuộc về Booking Session nào không'},
        {'type': 'flow', 'actor': '', 'system': 'Nếu thuộc Booking Session: AI phân tích trích xuất thông tin (giá, confirm)'},
        {'type': 'flow', 'actor': '', 'system': 'Hệ thống bắn notification (WebSocket) tới Frontend để hiển thị tin nhắn mới'},
        {'type': 'exception_header', 'label': 'Ngoại lệ', 'value': 'Lỗi phân tích tin nhắn: AI không hiểu nội dung (ví dụ voice, sticker), hệ thống hiển thị "Tin nhắn dạng media/không xác định"'},
    ])
    
    # Save document
    output_path = '/home/navin/Working/Booking_Agents/travel-agent-mvp/docs/usecases/zalo_usecases.docx'
    doc.save(output_path)
    print(f"✅ Document saved to: {output_path}")

if __name__ == "__main__":
    create_zalo_usecases_doc()
