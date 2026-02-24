from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re

def create_security_test_doc():
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)

    title = doc.add_paragraph('5.6. CHI TIẾT KỊCH BẢN KIỂM THỬ BẢO MẬT')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.bold = True
    title_run.font.size = Pt(16)
    doc.add_paragraph('')

    md_path = 'docs/detailed_security_testing.md'
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    sections = re.split(r'(### Bảng 5\.4\.\d+: [^\n]+)', content)
    
    for section in sections:
        section = section.strip()
        if not section: continue
        
        if section.startswith('### '):
            p = doc.add_paragraph(section.replace('### ', '').strip())
            p.style = 'Heading 3'
            p.italic = True
            
        elif '|' in section and 'STT' in section:
            lines = section.split('\n')
            rows = []
            for line in lines:
                if '|' in line and '---' not in line:
                    cols = [c.strip() for c in line.split('|') if c.strip()]
                    if cols:
                        rows.append(cols)
            
            if not rows: continue

            headers = rows[0]
            data_rows = rows[1:]
            
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = 'Table Grid'
            table.autofit = False 
            
            hdr_cells = table.rows[0].cells
            for i, h in enumerate(headers):
                p = hdr_cells[i].paragraphs[0]
                run = p.add_run(h)
                run.bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                tcPr = hdr_cells[i]._tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'E7E6E6')
                tcPr.append(shd)

            for r_data in data_rows:
                row = table.add_row()
                for i, val in enumerate(r_data):
                    if i < len(row.cells):
                        cell_text = val.replace('<br>', '\n').replace('**', '')
                        row.cells[i].text = cell_text
            
            table.columns[0].width = Inches(0.4)
            table.columns[1].width = Inches(1.5)
            table.columns[2].width = Inches(2.5)
            table.columns[3].width = Inches(2.0)
            
            doc.add_paragraph('')

    doc.save('docs/detailed_security_testing.docx')
    print("Created docs/detailed_security_testing.docx")

if __name__ == "__main__":
    create_security_test_doc()
