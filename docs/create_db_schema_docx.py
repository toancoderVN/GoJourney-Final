from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os
import re

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for edge in ('top', 'start', 'bottom', 'end', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = tcPr.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcPr.append(element)
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

def create_db_document():
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Title
    title = doc.add_paragraph('TỪ ĐIỂN DỮ LIỆU HỆ THỐNG (DATABASE SCHEMA)')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.bold = True
    title_run.font.size = Pt(18)
    title_run.font.color.rgb = RGBColor(44, 62, 80)
    doc.add_paragraph('')

    md_path = 'docs/FULL_DATABASE_SCHEMA.md'
    if not os.path.exists(md_path):
        print("Markdown file not found!")
        return

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    current_table = None
    table_rows = []
    
    for line in lines:
        line = line.strip()
        
        # Headers
        if line.startswith('## '):
            p = doc.add_paragraph(line.replace('#', '').strip())
            p.style = 'Heading 2'
        elif line.startswith('### '):
            p = doc.add_paragraph(line.replace('#', '').strip())
            p.style = 'Heading 3'
            # Reset table data
            table_rows = []
        
        # Table content
        elif line.startswith('|') and 'Field Name' not in line and '---' not in line:
            parts = [p.strip() for p in line.split('|') if p.strip()]
            if parts:
                table_rows.append(parts)
                
                # If next line is not a table row, render the table
                # (Simple logic: if we have rows and next line is empty or header)
                # But here we construct table immediately
                
    # Since parsing MD tables strictly is complex line-by-line, 
    # let's use a simpler known structure or regex.
    # Actually, let's hardcode the tables from the MD content I generated
    # This ensures perfect formatting.
    
    tables_data = {
        "User Service Domain": [
            ("user_profiles", [
                ("id", "uuid", "PK", "ID duy nhất"),
                ("email", "varchar", "Unique", "Email đăng nhập"),
                ("phone", "varchar", "Nullable", "SĐT"),
                ("firstName", "varchar", "Nullable", "Tên"),
                ("lastName", "varchar", "Nullable", "Họ"),
                ("preferences", "jsonb", "Nullable", "Cấu hình"),
                ("defaultBudgetMin", "decimal", "Nullable", "Min Budget"),
                ("defaultBudgetMax", "decimal", "Nullable", "Max Budget")
            ]),
            ("travel_companions", [
                ("id", "uuid", "PK", "ID"),
                ("userId", "uuid", "FK", "Owner"),
                ("companionId", "uuid", "FK", "Friend"),
                ("relationship", "enum", "Enum", "friend/family"),
                ("status", "enum", "Enum", "connected/pending")
            ])
        ],
        "Trip Service Domain": [
            ("trips", [
                ("id", "uuid", "PK", "ID chuyến đi"),
                ("userId", "uuid", "Index", "Người tạo"),
                ("name", "varchar", "NotNull", "Tên chuyến đi"),
                ("status", "enum", "Enum", "draft/confirmed"),
                ("budget", "jsonb", "Nullable", "Chi tiết ngân sách")
            ]),
            ("bookings", [
                ("id", "uuid", "PK", "ID booking"),
                ("tripId", "uuid", "FK", "Thuộc chuyến đi"),
                ("providerId", "uuid", "NotNull", "NCC"),
                ("status", "enum", "Enum", "pending/confirmed"),
                ("price", "jsonb", "NotNull", "Giá & Currency")
            ]),
            ("booking_conversations", [
                ("id", "uuid", "PK", "ID"),
                ("tripId", "uuid", "FK", "Trip"),
                ("messages", "jsonb", "NotNull", "Lịch sử chat Zalo")
            ])
        ]
    }
    
    # Render Tables
    for domain, tables in tables_data.items():
        doc.add_paragraph('')
        # Domain Header handled by MD logical parsing or separate add
        # Let's just create detailed tables for what we know is important

    # --- RE-IMPLEMENTING MD PARSER FOR GENERIC TABLES ---
    # This allows the script to work with whatever exists in FULL_DATABASE_SCHEMA.md
    
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    # Split by Heading 3 (Tables)
    sections = re.split(r'(### \d+\.\d+\. Bảng `[^`]+`)', content)
    
    header_processed = False
    
    for section in sections:
        section = section.strip()
        if not section: continue
        
        if section.startswith('### '):
            # Table Title
            p = doc.add_paragraph(section.replace('###', '').strip())
            p.style = 'Heading 3'
        elif section.startswith('Lưu trữ') or section.startswith('Danh sách') or section.startswith('Thực thể'):
            # Description text
            lines = section.split('\n')
            desc = lines[0] # First line is usually description
            doc.add_paragraph(desc).italic = True
            
            # Find markdown table
            table_lines = [l for l in lines if '|' in l]
            if len(table_lines) > 2:
                # Create Word Table
                # Filter outline chars
                rows_data = []
                headers = []
                
                for i, row in enumerate(table_lines):
                    if '---' in row: continue
                    cols = [c.strip().replace('**', '') for c in row.split('|') if c.strip()]
                    
                    if not cols: continue
                    
                    if not headers:
                        headers = cols
                    else:
                        rows_data.append(cols)
                
                if headers and rows_data:
                    table = doc.add_table(rows=1, cols=len(headers))
                    table.style = 'Table Grid'
                    
                    # Set Header
                    hdr_cells = table.rows[0].cells
                    for i, h in enumerate(headers):
                        if i < len(hdr_cells):
                            runner = hdr_cells[i].paragraphs[0].add_run(h)
                            runner.bold = True
                            runner.font.color.rgb = RGBColor(255, 255, 255)
                            
                            # Blue BG
                            tcPr = hdr_cells[i]._tc.get_or_add_tcPr()
                            shd = OxmlElement('w:shd')
                            shd.set(qn('w:val'), 'clear')
                            shd.set(qn('w:color'), 'auto')
                            shd.set(qn('w:fill'), '2E74B5')
                            tcPr.append(shd)

                    # Set Data
                    for row_data in rows_data:
                        row = table.add_row()
                        for i, val in enumerate(row_data):
                            if i < len(row.cells):
                                row.cells[i].text = val
                                set_cell_border(row.cells[i], bottom={"sz": 4, "val": "single", "color": "auto"})

    doc.save('docs/database_schema_design.docx')
    print("Created docs/database_schema_design.docx")

if __name__ == "__main__":
    create_db_document()
